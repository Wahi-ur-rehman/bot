from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_explanation_docx(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('ApexScalp AI v3.0: Technical Reference Manual', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "ApexScalp AI is a multi-layer, autonomous trading system. It utilizes a hybrid approach, combining "
        "traditional Technical Analysis (RSI, MACD, EMA), Price Action (Patterns), and Ensemble Machine Learning "
        "(Random Forest) to execute institutional-grade scalping trades on MetaTrader 5."
    )

    # Library Ecosystem
    doc.add_heading('2. Library Ecosystem & Rationale', level=1)
    doc.add_paragraph("The bot is built on a modern Python stack for low latency and high reliability.")
    
    libs = [
        ("MetaTrader5", "Core bridge to MT5 terminal. Handles data fetching and order execution."),
        ("scikit-learn", "AI/ML Engine. Provides Random Forest Classifier for predictive modeling."),
        ("pandas", "Data manipulation. Handles candle OHLCV dataframes and feature engineering."),
        ("numpy", "Numerical processing. Fast calculations for technical signals."),
        ("ta", "Technical Analysis library. Standardized implementations of indicators like ADX and RSI."),
        ("python-docx", "Documentation automated. Used to generate this technical manual.")
    ]
    
    lib_table = doc.add_table(rows=1, cols=2)
    lib_table.style = 'Table Grid'
    hdr = lib_table.rows[0].cells
    hdr[0].text = 'Library'
    hdr[1].text = 'Purpose & Rationale'
    
    for lib, desc in libs:
        row = lib_table.add_row().cells
        row[0].text = lib
        row[1].text = desc

    # Functional Blueprint
    doc.add_heading('3. Functional Blueprint (Core Operations)', level=1)
    doc.add_paragraph("This table maps every core trading operation to its specific function in the codebase.")
    
    funcs = [
        ("Main Execution Loop", "bot.run()", "Orchestrates the 60s scan cycle and coordinates all modules."),
        ("Data Retrieval", "mt5_connector.fetch_candles()", "Pulls live OHLCV data from the broker terminal."),
        ("Technical Features", "analysis_engine.compute_indicators()", "Calculates RSI, MACD, and EMA features."),
        ("Predictive Intelligence", "MLSignal.predict()", "Uses the Random Forest model to score trade probability."),
        ("Signal Confluence", "combined_signal()", "Weights all analysis layers into a single 'go/no-go' decision."),
        ("Volatility stops", "risk_manager.calculate_sl_tp()", "Calculates dynamic ATR-based profit and loss levels."),
        ("Position Sizing", "calculate_lot_size()", "Ensures max 0.5% equity risk per trade."),
        ("Execution", "trade_executor.place_order()", "Routes trade requests to the MetaTrader 5 deal server."),
        ("Broker Sync", "monitor_positions()", "Tracks TP/SL hits and detects manual closures by the broker.")
    ]
    
    func_table = doc.add_table(rows=1, cols=3)
    func_table.style = 'Table Grid'
    hdr = func_table.rows[0].cells
    hdr[0].text = 'Operation'
    hdr[1].text = 'Function Mapping'
    hdr[2].text = 'Technical Responsibility'
    
    for op, map, res in funcs:
        row = func_table.add_row().cells
        row[0].text = op
        row[1].text = map
        row[2].text = res

    # The Confluence logic
    doc.add_heading('4. The Confluence Logic', level=1)
    doc.add_paragraph("The bot uses a 'weighted consensus' model to filter noise:")
    doc.add_paragraph("Technical Indicators (35% Weight)", style='List Bullet')
    doc.add_paragraph("Price Action Patterns (25% Weight)", style='List Bullet')
    doc.add_paragraph("Machine Learning (40% Weight)", style='List Bullet')

    # Risk Protocol
    doc.add_heading('5. Risk Protocol', level=1)
    doc.add_paragraph("Risk Per Trade: 0.5% Fixed Fractional", style='List Bullet')
    doc.add_paragraph("Stop Loss Filter: 1.5x ATR (Dynamic Volatility)", style='List Bullet')
    doc.add_paragraph("Reward Ratio: 1:2 Minimum Target", style='List Bullet')
    doc.add_paragraph("Max Capacity: 8 Simultaneous Positions", style='List Bullet')

    doc.save(output_path)

if __name__ == "__main__":
    create_explanation_docx("ApexScalp_Technical_Manual_v3.docx")
