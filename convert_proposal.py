import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_to_txt(md_content, output_path):
    # Remove Mermaid diagrams (they don't look good in plain text)
    text = re.sub(r'```mermaid.*?```', '[Architecture/Workflow Diagram omitted in plain text]', md_content, flags=re.DOTALL)
    # Remove markdown formatting
    text = re.sub(r'#+\s+', '', text)
    text = re.sub(r'>\s*\[!.*?\]', 'IMPORTANT:', text)
    text = re.sub(r'---', '--------------------------------------------------', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def convert_to_docx(md_content, output_path):
    doc = Document()
    
    # Title
    title_match = re.search(r'^#\s+(.*)', md_content, re.MULTILINE)
    if title_match:
        h = doc.add_heading(title_match.group(1), 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle_match = re.search(r'^##\s+(.*)', md_content, re.MULTILINE)
    if subtitle_match:
        s = doc.add_paragraph(subtitle_match.group(1))
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.style = 'Subtitle'

    doc.add_paragraph('\n')
    
    # Process content
    lines = md_content.split('\n')
    in_mermaid = False
    in_table = False
    
    for line in lines:
        if line.startswith('# '): continue
        if line.startswith('## '): continue
        
        # Skip horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('__________________________________________________')
            continue
            
        # Mermaid skip
        if line.startswith('```mermaid'):
            in_mermaid = True
            doc.add_paragraph('[Architecture/Workflow Diagram included in original Markdown version]')
            continue
        if in_mermaid:
            if line.startswith('```'):
                in_mermaid = False
            continue
            
        # Callouts
        if line.startswith('> [!'):
            p = doc.add_paragraph()
            run = p.add_run("IMPORTANT: ")
            run.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0)
            continue
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph(text)
            p.style = 'Quote'
            continue

        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), 1)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), 2)
        
        # Tables (basic support)
        elif line.startswith('|'):
            if not in_table:
                in_table = True
                # Skip header separator line if it follows
            if ':---' in line:
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                p = doc.add_paragraph(" • " + " | ".join(cells))
        
        # Lists
        elif line.startswith('- '):
            in_table = False
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            in_table = False
            p = doc.add_paragraph(line[3:], style='List Number')
            
        # Standard paragraphs
        elif line.strip():
            in_table = False
            # Clean up bold and links
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text)
        else:
            in_table = False

    doc.save(output_path)

if __name__ == "__main__":
    with open("project_proposal.md", "r", encoding="utf-8") as f:
        content = f.read()
    
    convert_to_txt(content, "project_proposal.txt")
    print("Created project_proposal.txt")
    
    convert_to_docx(content, "project_proposal.docx")
    print("Created project_proposal.docx")
